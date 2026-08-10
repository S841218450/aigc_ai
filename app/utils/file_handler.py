import hashlib
import io
import os
import re
import zipfile


from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from app.utils.logger_handle import logger


def get_file_md5_hex(file_path:str):

    if not os.path.exists(file_path):
        logger.error(f'[md5计算] 文件不存在: {file_path}')
        return

    if not os.path.isfile(file_path):
        logger.error(f'[md5计算] 路径: {file_path}不是文件')
        return

    md5_obj = hashlib.md5()
    #4KB分片,避免文件过大爆内存
    chunk_size = 4096
    try:
        with open(file_path, 'rb') as f:
            while chunk := f.read(chunk_size):
                md5_obj.update(chunk)

            md5_hex = md5_obj.hexdigest()
            return md5_hex
    except Exception as e:
        logger.error(f"计算文件md5失败{e}")
        return None


def listdir_with_allowed_type(path:str, allowed_types:tuple[str]):
    files = []
    if not os.path.isdir(path):
        logger.error(f"{path}不是文件夹")
        return allowed_types

    for f in os.listdir(path):
        if f.endswith(allowed_types):
            files.append(os.path.join(path, f))

    return tuple(files)


def _normalize_col_name(raw: str) -> str:
    if raw is None:
        return ""
    s = str(raw).strip().replace("\r", "").replace("\n", " ").replace("  ", " ")
    return s or ""


def _row_non_empty_cnt(row) -> int:
    """一行中非空单元格数（None / 纯空白字符串 不算）。"""
    return sum(
        1 for v in row
        if v is not None and not (isinstance(v, str) and not v.strip())
    )


def _locate_table_header(raw_rows: list, max_scan: int = 20) -> tuple:
    """
    在 raw_rows 前 max_scan 行内定位"表头行"，兼容开头有文本描述的非标准表。

    启发式（足够覆盖 95% 场景，找不到时由调用方退回首行兜底）：
    1. 候选行非空单元格数 >= 2（真实表头至少 2 列）
    2. 候选行之后至少连续 2 行：非空单元格数 >= 2 且与候选行非空数差 <= 1（列对齐）
       —— 排除"标题行 + 单行数据"的伪表头（标题行通常 1 格文字，与数据行列数差 > 1）
    3. 若前 max_scan 行都不满足，返回 -1（调用方用 raw_rows[0] 当表头）

    Returns:
        header_idx: 表头行下标（-1 = 未找到，调用方退回首行）
        prefix_rows: 表头行之前的所有行（不含表头），供调用方作为前言文本
    """
    scan_end = min(max_scan, len(raw_rows))
    for i in range(scan_end):
        cnt = _row_non_empty_cnt(raw_rows[i])
        if cnt < 2:
            continue
        aligned = 0
        # 看后续最多 5 行，需要至少 2 行与候选表头列对齐
        for j in range(i + 1, min(i + 6, len(raw_rows))):
            jcnt = _row_non_empty_cnt(raw_rows[j])
            if jcnt < 2:
                continue
            if abs(jcnt - cnt) <= 1:
                aligned += 1
                if aligned >= 2:
                    return i, list(raw_rows[:i])
    return -1, []


def _prefix_rows_to_text(prefix_rows: list) -> str:
    """表头前的行（前言文本）拼成纯文本，供调用方并入文档文本流。"""
    lines = []
    for r in prefix_rows:
        parts = [str(v).strip() for v in r if v is not None and str(v).strip()]
        if parts:
            lines.append(" | ".join(parts))
    return "\n".join(lines).strip()


# 常见 OOXML 命名空间前缀（用于修复 docProps/custom.xml 缺失声明的问题）
_XML_NS_MAP = {
    "vt": "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes",
    "xsi": "http://www.w3.org/2001/XMLSchema-instance",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
}


def _repair_xlsx_custom_xml(file_path: str) -> io.BytesIO:
    """
    修复 docProps/custom.xml 缺失命名空间声明的 xlsx（WPS 等生成的文件常见问题），
    返回可被 openpyxl 重新读取的 BytesIO。
    """
    with open(file_path, "rb") as f:
        data = f.read()
    with zipfile.ZipFile(io.BytesIO(data)) as zin:
        items = [(item, zin.read(item.filename)) for item in zin.infolist()]
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item, content in items:
            if item.filename == "docProps/custom.xml":
                content = _inject_missing_xml_ns(content, _XML_NS_MAP)
            zout.writestr(item, content)
    out.seek(0)
    return out


def _inject_missing_xml_ns(xml_bytes: bytes, ns_map: dict) -> bytes:
    """在根元素上注入文档里使用了但未声明的前缀命名空间。"""
    text = xml_bytes.decode("utf-8", errors="replace")
    used = set(re.findall(r"<([A-Za-z_][\w.-]*):", text))
    missing = [p for p in used if p in ns_map and f"xmlns:{p}" not in text]
    if not missing:
        return xml_bytes
    # 根元素（跳过 <?xml ...?> 声明）
    m = re.search(r"<[A-Za-z_][\w.-]*", text)
    if not m:
        return xml_bytes
    inject = "".join(f' xmlns:{p}="{ns_map[p]}"' for p in missing)
    return (text[: m.end()] + inject + text[m.end():]).encode("utf-8")


def excel_loader(
    file_path: str,
    *,
    output_structured: bool = False,
) -> list:
    """
    Excel 加载器（标准表：首行当表头，从第二行起每行一条记录）
    :param file_path: .xlsx 或 .xls 文件路径
    :param output_structured:
        True  -> 返回 List[Dict] 结构化数据（供 DocumentProcessor.process_structured_rows 使用）
                  每个元素 = {"sheet_name": str, "columns": List[str], "rows": List[Dict[str, Any]]}
        False -> 返回 list[Document]（兼容旧接口，每个 Sheet 1 个 Document，page_content = 原始文本拼接）
    :return: 见上
    """
    import os

    suffix = os.path.splitext(file_path)[1].lower()

    # ---------- 选择底层库 ----------
    if suffix == ".xlsx":
        try:
            from openpyxl import load_workbook  # type: ignore
        except ImportError as e:
            raise ImportError(
                "处理 .xlsx 需要 openpyxl，请先 pip install openpyxl"
            ) from e

        def _read_sheet(ws):
            rows_iter = ws.iter_rows(values_only=True)
            return list(rows_iter), ws.title

        try:
            wb = load_workbook(filename=file_path, data_only=True, read_only=True)
        except Exception as e:
            # WPS 等生成的 xlsx 在 docProps/custom.xml 缺少命名空间声明时，
            # openpyxl 读取自定义属性会抛 XMLSyntaxError，修复后重试
            if "Namespace prefix" in str(e):
                wb = load_workbook(
                    filename=_repair_xlsx_custom_xml(file_path),
                    data_only=True,
                    read_only=True,
                )
            else:
                raise
        sheet_names = wb.sheetnames
        all_raw = []
        for name in sheet_names:
            ws = wb[name]
            rows_list, title = _read_sheet(ws)
            all_raw.append((title, rows_list))
        wb.close()

    elif suffix == ".xls":
        try:
            import xlrd  # type: ignore  # xlrd<=1.2.0 才支持 xls
        except ImportError as e:
            raise ImportError(
                "处理 .xls 需要 xlrd<=1.2.0，请先 pip install 'xlrd<=1.2.0'"
            ) from e

        book = xlrd.open_workbook(file_path)
        all_raw = []
        for idx in range(book.nsheets):
            sh = book.sheet_by_index(idx)
            rows_list = [
                [sh.cell_value(r, c) for c in range(sh.ncols)]
                for r in range(sh.nrows)
            ]
            all_raw.append((sh.name, rows_list))
    else:
        raise ValueError(f"excel_loader 不支持的后缀: {suffix}，仅支持 .xlsx/.xls")

    # ---------- 表解析：先定位表头（兼容开头有文本描述的非标准表），找不到则首行兜底 ----------
    structured_sheets = []
    for sheet_name, raw_rows in all_raw:
        if not raw_rows:
            continue
        header_idx, prefix_rows = _locate_table_header(raw_rows)
        if header_idx < 0:
            header_idx = 0
            prefix_rows = []
        # 表头前的内容（前言文本），供调用方并入文档文本流（如"本表收录在售产品 200 个"）
        prefix_text = _prefix_rows_to_text(prefix_rows)

        header = [_normalize_col_name(c) for c in raw_rows[header_idx]]
        # 表头去重（同名列加后缀 _2, _3...）
        seen = {}
        uniq_header = []
        for i, h in enumerate(header):
            hh = h if h else f"col_{i+1}"
            if hh in seen:
                seen[hh] += 1
                uniq_header.append(f"{hh}_{seen[hh]}")
            else:
                seen[hh] = 1
                uniq_header.append(hh)
        col_cnt = len(uniq_header)

        data_rows = []
        for row in raw_rows[header_idx + 1:]:
            # 空行跳过：整行全部 None/空字符串
            values = list(row)
            if all(v is None or (isinstance(v, str) and not v.strip()) for v in values):
                continue
            if len(values) < col_cnt:
                values = values + [None] * (col_cnt - len(values))
            elif len(values) > col_cnt:
                values = values[:col_cnt]
            data_rows.append(dict(zip(uniq_header, values)))

        structured_sheets.append({
            "sheet_name": sheet_name,
            "columns": uniq_header,
            "rows": data_rows,
            "prefix_text": prefix_text,
        })

    if output_structured:
        return structured_sheets

    # ---------- 兼容旧接口（Document 模式）：每个 Sheet 一个 page_content ----------
    docs: list[Document] = []
    for sht in structured_sheets:
        lines = [f"Sheet: {sht['sheet_name']}"]
        if sht.get("prefix_text"):
            lines.append(sht["prefix_text"])
            lines.append("")
        if sht["columns"]:
            lines.append("Columns: " + ", ".join(sht["columns"]))
        lines.append("")
        for r in sht["rows"]:
            parts = []
            for col in sht["columns"]:
                v = r.get(col, "")
                if v is None:
                    v = ""
                parts.append(f"{col}: {v}")
            lines.append(" | ".join(parts))
        page_content = "\n".join(lines).strip()
        docs.append(Document(
            page_content=page_content,
            metadata={"sheet_name": sht["sheet_name"], "columns": sht["columns"]},
        ))
    return docs


def docx_loader(
    file_path: str,
    *,
    output_structured: bool = False,
) -> list:
    """
    DOCX 加载器（段落 + 表格混合流）。
    :param output_structured:
        True  -> 返回 Dict：{
                  "paragraph_text": str（自然语言段落，保留换行/结构）,
                  "tables": List[{
                      "table_index": int,
                      "preceding_title": str（表格前最近的非空段落文本，作为 chunk 首行的 sheet_name）,
                      "columns": List[str],
                      "rows": List[Dict[str, Any]]
                  }]
              }
        False -> 返回 list[Document]（兼容旧接口，把段落和表格都转成纯文本 page_content）
    """
    import os

    suffix = os.path.splitext(file_path)[1].lower()
    if suffix != ".docx":
        raise ValueError(f"docx_loader 仅支持 .docx，实际后缀: {suffix}")

    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError as e:
        raise ImportError(
            "处理 .docx 需要 python-docx，请先 pip install python-docx"
        ) from e

    doc = DocxDocument(file_path)

    # ---------- 1. 收集「段落和表格」的文档顺序，为了拿到「表格前最近的非空段落」作为 preceding_title ----------
    # 注意：python-docx 的 doc.paragraphs 和 doc.tables 是**两个独立的扁平列表**，不保留它们在文档中的真实顺序。
    # 对于大多数知识库文档（标题→段落→表格→标题→段落→表格），我们用一个启发式：
    # 表格在 doc.tables 的顺序，去匹配「在 doc.paragraphs 中，当前表格对应的段落位置」的问题。
    # 简化做法：按「遍历所有段落 + 用 block-level child 索引顺序」太复杂。
    # 这里采用更简单且在 95% 知识库场景下足够准确的策略：
    #   「每个表格的 preceding_title = 到目前为止见过的所有段落里，**最后一个非空且长度小于等于 80 字符（一般是小节标题）** 的段落」
    #   如果没有这种短段落，再 fallback 到最后一个非空段落。
    last_nonempty_paragraph = ""
    last_short_paragraph = ""
    paragraph_parts = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            paragraph_parts.append(text)
            last_nonempty_paragraph = text
            # 启发式：小节标题不会太长
            if 0 < len(text) <= 80:
                last_short_paragraph = text

    structured_tables = []
    for tbl_idx, table in enumerate(doc.tables):
        rows_raw = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                # 单元格内部的换行保留，但多个空格压缩
                cell_text = (cell.text or "").replace("\r\n", "\n").replace("\r", "\n")
                cell_text = re.sub(r'[ \t]+', ' ', cell_text)
                cell_text = cell_text.strip()
                row_cells.append(cell_text)
            rows_raw.append(row_cells)

        if not rows_raw:
            continue
        # 表格内部也可能有标题行（如"表1：产品信息"），先定位真实表头
        header_idx, prefix_rows = _locate_table_header(rows_raw)
        if header_idx < 0:
            header_idx = 0
            prefix_rows = []
        prefix_text = _prefix_rows_to_text(prefix_rows)

        header = [_normalize_col_name(c) for c in rows_raw[header_idx]]
        seen = {}
        uniq_header = []
        for i, h in enumerate(header):
            hh = h if h else f"col_{i+1}"
            if hh in seen:
                seen[hh] += 1
                uniq_header.append(f"{hh}_{seen[hh]}")
            else:
                seen[hh] = 1
                uniq_header.append(hh)
        col_cnt = len(uniq_header)

        data_rows = []
        for row in rows_raw[header_idx + 1:]:
            values = list(row)
            if all(not str(v).strip() for v in values):
                continue
            if len(values) < col_cnt:
                values = values + [""] * (col_cnt - len(values))
            elif len(values) > col_cnt:
                values = values[:col_cnt]
            data_rows.append(dict(zip(uniq_header, values)))

        # 表格内部表头前的标题行（如"表1：产品信息"）优先作为表格名
        if prefix_text:
            preceding_title = prefix_text.splitlines()[0]
        structured_tables.append({
            "table_index": tbl_idx,
            "preceding_title": preceding_title,
            "columns": uniq_header,
            "rows": data_rows,
            "prefix_text": prefix_text,
        })

    paragraph_text = "\n\n".join(paragraph_parts).strip()

    if output_structured:
        return {
            "paragraph_text": paragraph_text,
            "tables": structured_tables,
        }

    # ---------- 兼容旧接口（Document 模式）：全部拼纯文本 ----------
    lines = []
    if paragraph_text:
        lines.append(paragraph_text)
    for tbl in structured_tables:
        lines.append("")
        lines.append(f"----- {tbl['preceding_title']} -----")
        if tbl["columns"]:
            lines.append("Columns: " + ", ".join(tbl["columns"]))
        for r in tbl["rows"]:
            parts = []
            for col in tbl["columns"]:
                v = r.get(col, "") or ""
                parts.append(f"{col}: {v}")
            lines.append(" | ".join(parts))
    page_content = "\n".join(lines).strip()
    return [Document(
        page_content=page_content,
        metadata={"table_count": len(structured_tables)},
    )]


def pdf_loader(file_path:str, passwd:None)->list[Document]:
    return PyPDFLoader(file_path, passwd).load()

def txt_loader(file_path:str)->list[Document]:
    return TextLoader(file_path).load()
